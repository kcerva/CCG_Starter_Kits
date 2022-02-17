#!/usr/bin/env python
# coding: utf-8

# In[2]:


import pandas as pd
import numpy as np
import plotly
import plotly_express as px
from plotly import tools
from plotly.offline import plot
import plotly.graph_objs as go
import plotly.io as pio
import os
from os import listdir
from os.path import isfile, join
from pathlib import Path
from os import listdir
from os.path import isfile, join
import country_converter as coco
import docx
# import accessdb
import pandas_access as mdb
from docx import Document
from docx.shared import Inches

#get_ipython().run_line_magic('matplotlib', 'inline')


# Setting working directory
os.chdir('C:\\Users\\KarlaC\\CCG_Starter_Kits\\')
if not os.path.exists('figures'):
    os.makedirs('figures')



# to do later, get country from xlsx file
 
country_data = pd.read_excel(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/Starter Kit - List of Countries.xlsx',  
                              )


# In[347]:


# get country list and turn into shor name and ISO
countries = list(country_data[country_data['Name of the Country'].notnull()]['Name of the Country'].unique())
short_names = coco.convert(names=countries, to='name_short')
iso3_codes = coco.convert(names=countries, to='ISO3', not_found=None)

ts_d = pd.DataFrame(list(zip(countries,short_names)), columns = ['Name of the Country','Country'])
ts_d = ts_d.set_index('Name of the Country')['Country'].to_dict()

ti3_d = pd.DataFrame(list(zip(countries,iso3_codes)), columns = ['Name of the Country','ISO3'])
ti3_d = ti3_d.set_index('Name of the Country')['ISO3'].to_dict()

country_data['ISO3'] = country_data['Name of the Country'].map(ti3_d)
country_data['Country'] = country_data['Name of the Country'].map(ts_d)

# Use this for running all available countries
# countries_short = list(country_data[country_data['Country'].notnull()]['Country'].unique())
# icountries = list(country_data[country_data['Country'].notnull()]['ISO3'].unique())

# Refinery capacity
refinery = country_data[['Country','ISO3',
                         'Refinery Capacity from McKinsery Refinery Reference Desk (enter no if 0, for sentence in article)']]
refinery = refinery.rename(columns= {'Refinery Capacity from McKinsery Refinery Reference Desk (enter no if 0, for sentence in article)':'RefCap'})
refinery = refinery[(refinery['RefCap']!='no')&(refinery['RefCap'].notnull())]

# Author contributions
authorcont = country_data[['Country','ISO3','Variable Author Contribution Statment']]
authorcont = authorcont.rename(columns = {'Variable Author Contribution Statment':'Authors'})
authorcont = authorcont[authorcont['Authors'].notnull()]




def melt_table(df):
    par = df.columns[1]
    df_pm = df
    df_pm.columns = df_pm.iloc[0]
    df_pm.columns = [df_pm.columns[0]]+[int(x) for x in df_pm.columns[1:]]
    df_pm = df_pm[1:]

    df_m = pd.melt(df_pm, 
                   id_vars = df_pm.columns[0], 
                   var_name = 'Year', value_name = 'Value')
    df_m['Parameter'] = par
    
    return df_m

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """
    output_para = output_doc_name.add_paragraph(style = 'List Number')
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
        
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

def melt_table_o(df):
    df_pm = df
    df_pm.columns = [str(x).replace('nan','Technology') for x in df.columns]
    df_m = pd.melt(df_pm, 
                   id_vars = df_pm.columns[0], 
                   var_name = 'Parameter', value_name = 'Value')
    return df_m


# In[349]:


# countries_short = country_data[country_data['Run_three']==1]['Country']
# icountries = country_data[country_data['Run_three']==1]['ISO3']

countries_short = ['Zimbabwe']
icountries = ['ZWE']

dib_names = ['Data in Brief Table 1',
             'Data in Brief Table 2',
             'Data in Brief Table 3 & Fig 1',
             'Data in Brief Table 4',
             'Data in Brief Table 5',
             'Data in Brief Table 6 & Fig 2',
             'Data in Brief Table 7',
             'Data in Brief Tables 8 & 9',
             'Data in Brief Figure 3']

dfn = ['cfile_3f','cfile_6f','cfile_9f']
labels = ['Capital Cost ($/kW)','Variable Costs ($/GJ)','Demand (PJ)']

sheets = ['AnnualElecProduction','TotalCapacityAnnual',
          'Transport', 'Annual CO2','ElecProductionByTS']
shortsheets = ['AnnualElecProduction','TotalCapacityAnnual',
          'Transport', 'Annual_CO2','ElecProductionByTS']
labels = ['PJ','GW','PJ','Mt CO2','PJ']
sk = [4,3,4,3,4]
restype = ['FF','LCv2','NZv2']
colors = px.colors.qualitative.Bold

sheets = ['AnnualElecProduction','TotalCapacityAnnual', 'Transport', 'Annual CO2']
shortsheets = ['AnnualElecProduction','TotalCapacityAnnual', 'Transport', 'Annual_CO2']
labels = ['PJ','GW','PJ','Mt CO2']
# dictionary for technology labels
tech_d = pd.read_excel(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/Data Preparation & Manipulation/Technology Codes.xlsx',
                                                  skiprows = 1)
col_d = pd.DataFrame(list(zip(tech_d['Code'],colors)))
tech_d = tech_d.set_index('Code')['Description'].to_dict()

tech_d["EMIC02"] = "CO2 Emissions"

    
# create figures
colors = px.colors.qualitative.Bold

for ic, country in zip(icountries,countries_short):
    if not os.path.exists('outputs/{}'.format(country)):
        os.makedirs('outputs/{}'.format(country))
    if not os.path.exists('figures/{}'.format(country)):
        os.makedirs('figures/{}'.format(country))
    authors = authorcont[authorcont['Country']==country]['Authors']

    # Read tables from preparation files
    for n, i in zip(dib_names, range(len(dib_names))):
        globals()['cfile_{}'.format(i+1)] = pd.read_excel(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/Data Preparation & Manipulation/New {} Data Collection.xlsx'.format(country),
                                                         sheet_name = n, na_values = ['#REF!'])


    # Input data figures

    # dfs for figures
    col0 = cfile_3.columns[0] 
    cfile_3f = cfile_3[~cfile_3[col0].isin(['No longer used','Coal Power Plant','Geothermal Power Plant',
                                          'Light Fuel Oil Power Plant','Oil Fired Gas Turbine (SCGT)',
                                           'Gas Power Plant (CCGT)', 'Crude Oil Refinery Option 1',
                                          'Crude Oil Refinery Option 2','Nuclear Power Plant',
                                           'Electricity Imports','Backstop Technology for ELC001',
                                          'Electricity Transmission', 'Electricity Distribution',
                                          'Light Fuel Oil Standalone Generator (1kW)', 'Electricity Exports'])]
    cfile_3f = cfile_3f.iloc[:,:8]
    cfile_3f.columns = cfile_3f.iloc[0]
    cfile_3f.columns = [cfile_3f.columns[0]]+[int(x) for x in cfile_3f.columns[1:]]
    cfile_3f = cfile_3f[1:]

    col0 = cfile_6.columns[0] 
    cfile_6f = cfile_6[~cfile_6[col0].isin(['Crude Oil Refinery Option 1',
                                          'Crude Oil Refinery Option 2'])]
    cfile_6f = cfile_6f.iloc[:,:8]
    cfile_6f.columns = cfile_6f.iloc[0]
    cfile_6f.columns = [cfile_6f.columns[0]]+[int(x) for x in cfile_6f.columns[1:]]
    cfile_6f = cfile_6f[1:]

    col0 = cfile_9.columns[0]
    cfile_9f = cfile_9[(cfile_9[col0].notnull())&(cfile_9[col0]!='Automaticaly linked to previous tabs')]
    cfile_9f.columns = cfile_9f.iloc[0]
    cfile_9f.columns = [cfile_9f.columns[0]]+[int(x) for x in cfile_9f.columns[1:]]
    cfile_9f = cfile_9f[1:]


    for s, l, f in zip(dfn, labels, range(3)):
        if s == 'cfile_9f':
            data = globals()[s][globals()[s]['Demand']=='Total electricity demand']
            
        else:
            data = globals()[s]
        fig = go.Figure()
        cone = data.columns[0]
        for i, c in zip(data[cone].unique(), colors): 
            fig.add_trace(go.Scatter(x = data.columns[1:], 
                                     y = data[data[cone]==i].values[0][1:],
                                     mode = 'lines',
                                     name = i ,
                                     line_color = c, 
                                         ))

        fig.update_layout( yaxis_title =  l, 
                            template = 'simple_white+presentation',

                         )



        pio.write_image(fig, r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig{1}_{2}.jpg'.format(country,f+1,ic), width = 1200, height = 700)
#         fig.show()

# Results figures

# Reading in results
    for r in restype:
        for s,sh, k in zip(sheets,shortsheets, sk):
    #         globals()[s] = pd.read_excel(path+'Results - {} BAUv10.xlsx'.format(country),
    #                          sheet_name = s, skiprows = k)
            globals()['{0}_{1}'.format(r,sh)] = pd.read_excel(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/Results/Excel Results Files/Checked/{0} {1} Results.xlsx'.format(country,r),
                                                             sheet_name = s, skiprows = k)
            globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)][globals()['{0}_{1}'.format(r,sh)][globals()['{0}_{1}'.format(r,sh)].columns[0]]!='Grand Total']

            globals()['{0}_{1}'.format(r,sh)].columns = ['Year']+list(globals()['{0}_{1}'.format(r,sh)].columns[1:])
            globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)].drop(columns = 'Grand Total')
#             if r == 'FF':
#                 Primary_Fuel_Production = pd.read_excel(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/Results/Excel Results Files/Checked/{0} {1} Results.xlsx'.format(country,r),
#                                          sheet_name = 'Primary Fuel Production', skiprows = 4)
#                 Primary_Fuel_Production = Primary_Fuel_Production[Primary_Fuel_Production[Primary_Fuel_Production.columns[0]]!='Grand Total']
#                 Primary_Fuel_Production.columns = ['Year']+list(Primary_Fuel_Production.columns[1:])
#                 Primary_Fuel_Production = Primary_Fuel_Production.drop(columns = 'Grand Total')
                
            if s == 'Annual CO2':
                mtname = [x for x in globals()['{0}_{1}'.format(r,sh)].columns if 'Mt' in x][0]
                globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)].drop(columns = ['Unnamed: 3','EMIC02'])
                globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)].rename(columns ={mtname:'EMIC02'})
                if 'Unnamed: 5' in globals()['{0}_{1}'.format(r,sh)].columns:
                    globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)].drop(columns = ['Unnamed: 5'])
            if s == 'ElecProductionByTS':
                # Adding column for year, filling it in and removing the slice totals per year 
                globals()['{0}_{1}'.format(r,sh)]['Slice'] = globals()['{0}_{1}'.format(r,sh)]['Year']
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)]['Slice']=='2020','Year'] = 2020
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)]['Slice']=='2030','Year'] = 2030
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)]['Slice']=='2040','Year'] = 2040
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)]['Slice']=='2050','Year'] = 2050
                
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)].loc[1:96].index,'Year'] = 2020
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)].loc[97:193].index,'Year'] = 2030
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)].loc[194:290].index,'Year'] = 2040
                globals()['{0}_{1}'.format(r,sh)].loc[globals()['{0}_{1}'.format(r,sh)].loc[291:].index,'Year'] = 2050
                
                globals()['{0}_{1}'.format(r,sh)] = globals()['{0}_{1}'.format(r,sh)][~globals()['{0}_{1}'.format(r,sh)]['Slice'].isin(['2020','2030','2040','2050'])] 
                

# Results figures

# colors = px.colors.qualitative.Bold

# for ic, country in zip(icountries,countries_short):
    for r in restype:
        for s, sh, l in zip(sheets, shortsheets, labels):
            data = globals()['{0}_{1}'.format(r,sh)]
            fig = go.Figure()
            if s == 'Annual CO2':
                fig.add_trace(go.Scatter(y = data['EMIC02'], 
                                         x = data['Year'],
                                         mode = 'lines',
                                         name = tech_d[i],
                                         line_color = colors[0], 
                                         )
                              )
                fig.update_layout( yaxis_title =  l, 
                                    template = 'simple_white+presentation',
                                 )
                pio.write_image(fig, r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_{1}{2}_{3}.jpg'.format(country,r,sh,ic), width = 1000, height = 600)

            else:
                for i, c in zip(data.columns[1:], colors): 
                    fig.add_trace(go.Bar(y = data[i], 
                                         x = data['Year'],
                                         name = tech_d[i],
                                         marker_color = c, 
                                         )
                                 )
    #             if s == 'ElecProductionByTS':
    #                 cols = [x for x in data.columns if (x!= 'Slice' and x!='Year')]
    #                 for y in data['Year'].unique():
    #                     for i, c in zip(data['Slice'].unique(), colors):
    #                         fig.add_trace(go.Bar(x = cols, 
    #                                          y = data[data['Year']==y]['Slice'],
    #                                          mode = 'lines',
    #                                          name = tech_d[i] ,
    #                                          line_color = c, 
    #         #                                  orientation = 'v',
    #                                          )
    #                                  )
                fig.update_layout( yaxis_title =  l, 
                                    template = 'simple_white+presentation',
                                  barmode='stack'
                                 )


                pio.write_image(fig, r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_{1}{2}_{3}.jpg'.format(country,r,sh,ic), width = 1000, height = 600)


# Making DiB documents

    # loop iso and country
# for ic, country in zip(icountries, countries_short):
    
    # Variable Reference Imports
    refl = country_data[country_data['Country']==country]['Reference List'].values[0]
    input_doc = Document(r'C:/Users/KarlaC/Loughborough University/CCG - SHARED- CCG partners/OA4_Platform/Starter Kit/PAPERS/Data in Brief Template/Reference List {}.docx'.format(refl))
    # output_doc = Document(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/DiB_{1}.docx'.format(country,ic))

    document = Document()

    document.add_heading("Selected ‘Starter Kit’ energy system modelling data for {} (#CCG)".format(country), 0)
    document.add_heading('Authors', 1)
    lauth = country_data[country_data['Country']==country]['Author List (draft, to be checked)'].values[0]
    document.add_paragraph('{}'.format(lauth))

    document.add_heading('Affiliations', 1)
    laff = country_data[country_data['Country']==country]['Author Affiliations'].values[0]
    document.add_paragraph('{}'.format(laff))

    document.add_heading('Corresponding author(s)', 1)
    document.add_paragraph('Lucy Allington (lallington9@gmail.com) and Carla Cannone (c.cannone@lboro.ac.uk)')

    document.add_heading('Abstract', 1)
    document.add_paragraph('Energy system modelling can be used to assess the implications of different scenarios and support improved policymaking. However, access to data is often a barrier to starting energy system modelling in developing countries, thereby causing delays. Therefore, this article provides data that can be used to create a simple zero order energy system model for {}, which can act as a starting point for further model development and scenario analysis. The data are collected entirely from publicly available and accessible sources, including the websites and databases of international organizations, journal articles, and existing modelling studies. This means that the dataset can be easily updated based on the latest available or more detailed local data. These data were also used to calibrate a simple energy system model using the Open Source Energy Modelling System (OSeMOSYS) and three stylized scenarios (Fossil Future, Least Cost and Net Zero by 2050) for 2020-2050. The assumptions used and results of these scenarios are presented in the appendix as an illustrative example of what can be done with these data. This simple model can be adapted and further developed by in-country analysts and academics, providing a platform for future work.'.format(country))

    document.add_heading('Keywords', 1)
    document.add_paragraph('U4RIA, Renewable energy, Cost-optimization, {}, Energy policy, CCG, OSeMOSYS'.format(country))

    document.add_heading('Specifications Table', 1)
    records = (
    #     ('Subject', 'Energy'),
        ('Specific subject area', 'Energy System Modelling'),
        ('Type of data', 'Tables\n Graphs\n Charts\n Description of modelling assumptions'),
        ('How data were acquired', 'Literature survey (databases and reports from international organisations; journal articles)'),
        ('Data format', 'Raw and Analysed'),
        ('Parameters for data collection', 'Data collected based on inputs required to create an energy system model for {}'.format(country)),
        ('Description of data collection', 'Data were collected from the websites, annual reports and databases of international organisations, as well as from academic articles and existing modelling databases.'),
        ('Data source location', 'Not applicable'),
        ('Data accessibility', 'With the article'),
    #     ('Related research article', 'Author’s name, Title, Journal, DOI/In Press'),
    )
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Subject'
    hdr_cells[1].text = 'Energy'
    for s1, s2 in records:
        row_cells = table.add_row().cells
        row_cells[0].text = s1
        row_cells[1].text = s2

    document.add_heading('Value of the data', 1)
    document.add_paragraph('These data can be used to develop national energy system models to inform national energy investment outlooks and policy plans, as well as provide insights on the evolution of the electricity supply system under different trajectories. ', style = 'List Bullet')
    document.add_paragraph('The data are useful for country analysts, policy makers and the broader scientific community, as a zero-order starting point for model development.', style = 'List Bullet')
    document.add_paragraph("These data could be used to examine a range of possible energy system pathways, in addition to the examples given in this study, to provide further insights on the evolution of the country's power system.",style = 'List Bullet')
    document.add_paragraph('The data can be used both for conducting an analysis of the power system but also for capacity building activities. Also, the methodology of translating the input data into modelling assumptions for a cost-optimization tool is presented here which is useful for developing a zero order Tier 2 national energy model [1]. This is consistent with U4RIA1 energy planning goals [2]. ',style = 'List Bullet')

    document.add_heading('1 Data Description', 1)
    document.add_paragraph('The data provided in this paper can be used as input data to develop an energy system model for {}. As an illustration, these data were used to develop an energy system model using the cost-optimization tool OSeMOSYS for the period 2015-2050. For reference, that model is described in Appendix A and its datafiles are available as Supplementary Materials. For reference that model and its datafile derived from this paper are available in Appendix 2. The data provided were collected from publicly available sources, including the reports of international organizations, journal articles and existing model databases. It includes the techno-economic parameters of supply-side technologies, installed capacities, emissions factors and power demands.'.format(country))

    records2 = (
    #     ('Item', 'Description of Content'),
        ('Table 1', 'A table showing the estimated installed capacity of different power plant types in {} for 2015-2018 '.format(country)),
        ('Table 2', 'A table showing cost and performance parameters for power generation technologies '),
        ('Table 3', 'A table showing capital cost projections for renewable energy technologies up to 2070 '),
        ('Figure 1', 'A graph showing capital cost projections for renewable energy technologies from 2015-2050'),
        ('Table 4', 'A table showing cost and performance parameters for power transmission and distribution technologies'),
        ('Table 5', 'A table showing cost and performance data for refinery technologies'),
        ('Table 6', 'A table showing fuel price projections up to 2070'),
        ('Figure 2', 'A graph showing fuel price projections from 2015-2050 '),
        ('Table 7', 'A table showing carbon dioxide emissions factors by fuel'),
        ('Table 8', 'A table showing estimated renewable energy potential in {}'.format(country)),
        ('Table 9', 'A table showing estimated fossil fuel reserves in {}'.format(country)),
        ('Figure 3', 'A graph showing an electricity demand projection for {} from 2015-2070'.format(country))
    )
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Description of Content'
    for s1, s2 in records2:
        row_cells = table.add_row().cells
        row_cells[0].text = s1
        row_cells[1].text = s2

    # document.add_heading('Main modelling assumptions', 2)

    document.add_heading('1.1 Existing Electricity Supply System', 2)

    # removing rows that aren't used anymore
    col0 = cfile_1.columns[0]
    cfile_1 = cfile_1[~cfile_1[col0].isin(['No longer used'])]
    cfile_1 = cfile_1.iloc[:,:5] # removing columns without importance
    # Removing zeros
    colsa = cfile_1.columns[1:]
    cfile_1 = cfile_1.loc[(cfile_1[colsa] > 0).all(axis=1)]
    # TO DO: remove rows where values are zero df = df[(df.T != 0).any()]
    cfile_1 = cfile_1.round(2) # rounding values
    cfile_1s = cfile_1.applymap(str)

    cmo = cfile_1.columns[-1]
    sumcap = cfile_1[cfile_1[cmo]!=2018][cmo].sum().round(2) #getting this sum from file
    document.add_paragraph('The total power generation capacity in {0} is estimated at {1} MW in 2018 [3,4,5,6]. The estimated existing power generation capacity is detailed in Table 1 below [3,4,5,6]. The methods used to calculate these estimates are described in more detail in Section 2.1.'.format(country,str(sumcap)))

    document.add_paragraph('Table 1: Installed Power Plants Capacity in {} [3,4,5,6]'.format(country))
    # Subtle Reference?
    table = document.add_table(rows=1, cols=5)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Estimated Installed Capacity (MW)'
    hdr_cells[2].text = ''
    hdr_cells[3].text = ''
    hdr_cells[4].text = ''

    h2 = ['Power Generation Technology', '2015','2016','2017','2018']
    row_cells = table.add_row().cells
    row_cells[0].text = h2[0]
    row_cells[1].text = h2[1]
    row_cells[2].text = h2[2]
    row_cells[3].text = h2[3]
    row_cells[4].text = h2[4] 

    # hdr_cells[0].text = 'Power Generation Technology'
    # hdr_cells[1].text = '2015'
    # hdr_cells[2].text = '2016'
    # hdr_cells[3].text = '2017'
    # hdr_cells[4].text = '2018'
    for ind, row in cfile_1s[1:].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]

    # Making csv
    table1 = melt_table(cfile_1)
    table1.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table1_{1}.csv'.format(country,ic), index = False)

    document.add_heading('1.2 Techno-economic Data for Power Generation Technologies', 2)
    document.add_paragraph('The techno economic parameters of power generation technologies are presented in Table 2, including costs, operational lives, efficiencies and average capacity factors. Cost, efficiency and operational life data were collected from reports by the International Renewable Energy Agency [7,8,9]. These cost data include projected cost reductions for renewable energy technologies, which are presented in Table 3. Country-specific capacity factors for solar PV, wind and hydropower technologies were sourced from Renewables Ninja and the PLEXOS-World 2015 Model Dataset [3,10,11], while generic capacity factors for other technologies were sourced from the International Renewable Energy Agency [8,12]. In the table below, average capacity factors across the whole year are presented, however these are day-time averages for solar PV technologies.')

    # reduce number of decimals
    document.add_paragraph('Table 2: Techno-economic parameters of power generation technologies [3,7,8,9,10,11,12]')
    table = document.add_table(rows=1, cols=6)
    table.style = 'TableGrid'
    # Removing columns with 'source' in name
    colt2 = [x for x in cfile_2.columns if not ('Source' in x or 'source' in x)]
    colt2 = [x for x in colt2 if x!='Variable Cost ($/GJ in 2020)']
    cfile_2 = cfile_2[colt2]
    # need to remove mmore rows
    col0 = cfile_2.columns[0]
    cfile_2 = cfile_2[~cfile_2[col0].isin(['No longer used','Electricity Imports',
                                           'Electricity Transmission','Electricity Distribution',
                                           'Backstop Technology for ELC001','Electricity Exports'])]
    cfile_2 = cfile_2.iloc[:,:6].round(2) # rounding values
    cfile_2s = cfile_2.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cfile_2.columns[0]
    hdr_cells[1].text = cfile_2.columns[1]
    hdr_cells[2].text = cfile_2.columns[2]
    hdr_cells[3].text = cfile_2.columns[3]
    hdr_cells[4].text = cfile_2.columns[4]
    hdr_cells[5].text = cfile_2.columns[5]
    for ind, row in cfile_2s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
        row_cells[5].text = row[5]
    # Making csv
    table2 = melt_table_o(cfile_2)
    table2.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table2_{1}.csv'.format(country,ic), index = False)

    document.add_paragraph()

    # fix which table it takes
    document.add_paragraph('Table 3: Projected costs of renewable power generation technologies for selected years to 2050. [7,9]')
    table = document.add_table(rows=1, cols=7)
    table.style = 'TableGrid'
    col0 = cfile_3.columns[0] 
    cfile_3 = cfile_3[~cfile_3[col0].isin(['No longer used','Coal Power Plant','Geothermal Power Plant',
                                          'Light Fuel Oil Power Plant','Oil Fired Gas Turbine (SCGT)',
                                           'Gas Power Plant (CCGT)', 'Crude Oil Refinery Option 1',
                                          'Crude Oil Refinery Option 2','Nuclear Power Plant',
                                           'Electricity Imports','Backstop Technology for ELC001',
                                          'Electricity Transmission', 'Electricity Distribution',
                                          'Light Fuel Oil Standalone Generator (1kW)', 'Electricity Exports'])]
    cfile_3 = cfile_3.round(2) # rounding values
    cfile_3 = cfile_3.iloc[:,:7]
    cfile_3s = cfile_3[1:]
    cfile_3s = cfile_3s.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Capital Cost ($/kW)'
    hdr_cells[2].text = ''
    hdr_cells[3].text = ''
    hdr_cells[4].text = ''
    hdr_cells[5].text = ''
    hdr_cells[6].text = ''

    # hdr_cells[0].text = 'Technology'
    # hdr_cells[1].text = '2015'
    # hdr_cells[2].text = '2020'
    # hdr_cells[3].text = '2025'
    # hdr_cells[4].text = '2030'
    # hdr_cells[5].text = '2040'
    # hdr_cells[6].text = '2050'
    h2 = ['Power Generation Technology', '2015','2020','2025','2030','2040','2050']
    # for s,si in zip(h2, range(len(h2))):
    row_cells = table.add_row().cells
    row_cells[0].text = h2[0]
    row_cells[1].text = h2[1]
    row_cells[2].text = h2[2]
    row_cells[3].text = h2[3]
    row_cells[4].text = h2[4] 
    row_cells[5].text = h2[5]
    row_cells[6].text = h2[6]

    for ind, row in cfile_3s.iloc[:,:7].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
        row_cells[5].text = row[5]
        row_cells[6].text = row[6]
    # Making csv
    table3 = melt_table(cfile_3)
    table3.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table3_{1}.csv'.format(country,ic), index = False)

    # Add figure
    document.add_paragraph('Figure 1: Projected costs of renewable power generation technologies for selected years to 2050 [7,9] ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig1_{1}.jpg'.format(country,ic), width=Inches(6))

    document.add_heading('1.3 Techno-economic Data for Power Transmission and Distribution', 2)
    col0 = cfile_4.columns[0]
    cfile_4 = cfile_4[cfile_4[col0]!='Reference']
    cols = [x for x in cfile_4.columns if ((x!= 'Fixed Cost ($/kW/yr in 2020)') and ( x != 'Variable Cost ($/GJ in 2020)'))]
    cfile_4 = cfile_4[cols]
    cfile_4 = cfile_4.iloc[:,:6]
    cfile_4[cfile_4.columns[-1]] = [np.round(x,2) for x in cfile_4[cfile_4.columns[-1]]]
    cfile_4[cfile_4.columns[-2]] = [np.round(x,2) for x in cfile_4[cfile_4.columns[-2]]]
    cfile_4[cfile_4.columns[-3]] = [np.round(x,2) for x in cfile_4[cfile_4.columns[-3]]]
    cfile_4 = cfile_4.round(2) # rounding values
    cfile_4s = cfile_4.applymap(str)
    # x1 = 15 #get from tables
    x1 = np.round(cfile_4['Efficiency (2030)'][0]*100,2)
    # x2 = 20 #get from tables
    x2 = np.round(cfile_4['Efficiency (2030)'][1]*100,2)
    document.add_paragraph('{0} has relatively high transmission and distribution losses. The techno-economic parameters of transmission and distribution technologies were taken from the Reference Case scenario of The Electricity Model Base for Africa (TEMBA) [13]. According to these data, the efficiencies of power transmission and distribution in {0} are expected to reach {1}% and {2}% respectively in 2030. In the following table, the techno-economic parameters associated with the transmission and distribution network are presented.'.format(country,str(x1),str(x2)))

    # fix which table it takes
    document.add_paragraph('Table 4: Techno-economic parameters for transmission and distribution technologies [13]')
    table = document.add_table(rows=1, cols=6)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cfile_4.columns[0]
    hdr_cells[1].text = cfile_4.columns[1]
    hdr_cells[2].text = cfile_4.columns[2]
    hdr_cells[3].text = cfile_4.columns[3]
    hdr_cells[4].text = cfile_4.columns[4]
    hdr_cells[5].text = cfile_4.columns[5]
    for ind, row in cfile_4s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
        row_cells[5].text = row[5]

    # Making csv
    table4 = melt_table_o(cfile_4)
    table4.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table4_{1}.csv'.format(country,ic), index = False)

    document.add_heading('1.4 Techno-economic Data for Refineries', 2)
    # dr = 30
    dr = refinery[refinery['Country']==country]['RefCap']
    if dr.empty:
        dr = 0
    document.add_paragraph('{0} has an estimated {1} domestic refinery capacity [14]. In the OSeMOSYS model, two oil refinery technologies were made available for investment in the future, each with different output activity ratios for Heavy Fuel Oil (HFO) and Light Fuel Oil (LFO). The technoeconomic data for these technologies are shown in Table 5.'.format(country,str(dr)))
    # fix which table it takes
    document.add_paragraph('Table 5: Techno-economic parameters for refinery technologies [14,15]')
    table = document.add_table(rows=1, cols=5)
    table.style = 'TableGrid'
    col0 = cfile_5.columns[0]
    cfile_5 = cfile_5[cfile_5[col0]!='Reference']
    cols = [x for x in cfile_5.columns if x != 'Efficiency (2020)']
    cfile_5 = cfile_5[cols]
    cfile_5 = cfile_5.iloc[:,:5].round(2) # rounding values
    cfile_5s = cfile_5.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cfile_5.columns[0]
    hdr_cells[1].text = cfile_5.columns[1]
    hdr_cells[2].text = cfile_5.columns[2]
    hdr_cells[3].text = cfile_5.columns[3]
    hdr_cells[4].text = cfile_5.columns[4]
    for ind, row in cfile_5s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
    # Making csv
    table5 = melt_table_o(cfile_5)
    table5.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table5_{1}.csv'.format(country,ic), index = False)

    document.add_heading('1.5 Fuel Prices', 2)
    document.add_paragraph('Estimated costs are provided for both imported and domestically-extracted fuels. The fuel price projections until 2050 are presented below. These are generic estimates based on an international oil price forecast [16] and cost estimates for Africa [8]. A detailed explanation of how these estimates were calculated is provided in section 2.2.')

    # reduce number of decimals
    document.add_paragraph('Table 6: Fuel price projections to 2060 [16,8]')
    table = document.add_table(rows=1, cols=7)
    table.style = 'TableGrid'
    col0 = cfile_6.columns[0] 
    cfile_6 = cfile_6[~cfile_6[col0].isin(['Crude Oil Refinery Option 1',
                                          'Crude Oil Refinery Option 2'])]
    cfile_6 = cfile_6.round(2) # rounding values
    cfile_6 = cfile_6.iloc[:,:8]
    cfile_6 = cfile_6[1:]
    cfile_6s = cfile_6.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'Fuel Price ($/GJ) '
    hdr_cells[2].text = ''
    hdr_cells[3].text = ''
    hdr_cells[4].text = ''
    hdr_cells[5].text = ''
    hdr_cells[6].text = ''

    h2 = ['Commodity', '2015','2020','2025','2030','2040','2050']
    row_cells = table.add_row().cells
    row_cells[0].text = h2[0]
    row_cells[1].text = h2[1]
    row_cells[2].text = h2[2]
    row_cells[3].text = h2[3]
    row_cells[4].text = h2[4] 
    row_cells[5].text = h2[5]
    row_cells[6].text = h2[6]

    # hdr_cells[0].text = 'Commodity'
    # hdr_cells[1].text = '2015'
    # hdr_cells[2].text = '2020'
    # hdr_cells[3].text = '2025'
    # hdr_cells[4].text = '2030'
    # hdr_cells[5].text = '2040'
    # hdr_cells[6].text = '2050'
    for ind, row in cfile_6s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
        row_cells[2].text = row[2]
        row_cells[3].text = row[3]
        row_cells[4].text = row[4]
        row_cells[5].text = row[5]
        row_cells[6].text = row[6]
    # Making csv
    table6 = melt_table(cfile_6)
    table6.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table6_{1}.csv'.format(country,ic), index = False)

    # Add figure
    document.add_paragraph('Figure 2: Fuel price projections to 2050 [16,8]')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig2_{1}.jpg'.format(country,ic), width=Inches(6))

    document.add_heading('1.6 Emission Factors', 2)
    document.add_paragraph('Fossil fuel technologies emit several gases, including carbon dioxide, methane and nitrous oxides throughout their operational lifetime. In this analysis, only carbon dioxide emissions are considered. These are accounted for using carbon dioxide emission factors assigned to each fuel, rather than each power generation technology. The assumed emission factors are presented in Table 7.')

    document.add_paragraph('Table 7: Fuel-specific CO2 Emission Factors [17]')
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'

    col0 = cfile_7.columns[0]
    cfile_7 = cfile_7[cfile_7[col0]!= 'Reference']
    col0 = cfile_7.columns[0]
    cfile_7 = cfile_7[cfile_7[col0].notnull()]
    cfile_7 = cfile_7.round(2) # rounding values
    cfile_7 = cfile_7.iloc[:,:2]
    cfile_7s = cfile_7.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = cfile_7.columns[0]
    hdr_cells[1].text = 'CO2 Emission Factor (kg CO2/GJ)'
    for ind, row in cfile_7s.iloc[:,:2].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    # Making csv
    table7 = melt_table_o(cfile_7)
    table7.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table7_{1}.csv'.format(country,ic), index = False)

    document.add_heading('1.7 Renewable and Fossil Fuel Reserves', 2)
    document.add_paragraph('Tables 8 and 9 show estimated domestic renewable energy potentials and fossil fuel reserves respectively in {}.'.format(country))

    # reduce number of decimals
    document.add_paragraph('Table 8: Estimated Renewable Energy Potentials [18,19,20]')
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'

    cfile_8.columns = cfile_8.iloc[0]
    cfile_8 = cfile_8[1:]
    cfile_88 = cfile_8.iloc[:,:2]
    col0 = cfile_88.columns[1]
    cfile_88 = cfile_88[cfile_88[col0].notnull()]
    cfile_88 = cfile_88.round(2) # rounding values
    cfile_8s = cfile_88.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = cfile_88.columns[1]
    for ind, row in cfile_8s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    # Making csv
    table88 = melt_table_o(cfile_88)
    table88.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table8_{1}.csv'.format(country,ic), index = False)

    document.add_paragraph('Table 9: Estimated Fossil Fuel Reserves [21,22] ')
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = ''
    # hdr_cells[1].text = 'Installed Capacity (GW)'
    # hdr_cells[2].text = ''
    cfile_89 = cfile_8.iloc[:,3:5]
    col0 = cfile_89.columns[0]
    cfile_89 = cfile_89[(cfile_89[col0].notnull())&                        (cfile_89[col0]!='(probably exclude the fossil table for Kenya since the sources state 0 proven reserves)')]
    col1 = cfile_89.columns[1]
    cfile_89[col1].replace(np.nan,0, inplace = True)
    cfile_89 = cfile_89.round(2) # rounding values
    cfile_8s = cfile_89.applymap(str)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = cfile_89.columns[1]
    for ind, row in cfile_8s.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]
    # Making csv
    table89 = melt_table_o(cfile_89)
    table89.to_csv(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/Table9_{1}.csv'.format(country,ic), index = False)

    document.add_heading('1.8 Electricity Demand Projection ', 2)
    col0 = cfile_9.columns[0]
    cfile_9e = cfile_9[(cfile_9[col0].notnull())&(cfile_9[col0]!='Automaticaly linked to previous tabs')]
    cfile_9e.columns = cfile_9e.iloc[0]
    cfile_9e.columns = [cfile_9e.columns[0]]+[int(x) for x in cfile_9e.columns[1:]]
    cfile_9e = cfile_9e[1:]
    cfile_9e = cfile_9e[cfile_9e['Demand']=='Total electricity demand']

    # e1 = 250 #find in table
    e1 = np.round(cfile_9e[2018].values[0],2)
    e2 = 500 #find in table
    e2 = np.round(cfile_9e[2030].values[0],2)
    # e3 = 200 #find
    e3 = np.round(cfile_9e[2050].values[0],2)
    document.add_paragraph('Electricity demand in {0} was estimated at {1} PJ in 2018 and is forecasted to reach {2} PJ by 2030 and {3} PJ by 2050 [23] in a reference scenario. Figure 3 below shows the electricity demand projection.'.format(country,str(e1),str(e2),str(e3)))

    # Add figure
    document.add_paragraph('Figure 3: Electricity Demand Projection (PJ) [23] ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig3_{1}.jpg'.format(country,ic), width=Inches(6))

    document.add_heading('2 Experimental Design, Materials, and Methods', 2)
    document.add_paragraph('Data were primarily collected from the reports and websites of international organizations, including the International Renewable Energy Agency (IRENA), the International Energy Agency (IEA)/UN Stats, the Asia-Pacific Economic Cooperation (APEC) and the Intergovernmental Panel on Climate Change (IPCC). Additionally, data were sourced from The Electricity Model Base for Africa (TEMBA), an existing OSeMOSYS model of African electricity supply.')

    document.add_heading('2.1 Electricity Supply System Data', 2)
    esupply = document.add_paragraph("Data on {}'s existing on-grid power generation capacity, presented in Table 1, were extracted from the PLEXOS World dataset [3,4,5] using scripts from OSeMOSYS global model generator [24]. ".format(country)) 
    esupply.add_run("PLEXOS World provides data for 2015, including the capacity and commissioning date of each power plant. These data were used to estimate installed capacity in future years based on the operational life data in Table 2. Data on {}'s off-grid capacity were sourced from yearly capacity statistics produced by IRENA [5]. Cost, efficiency and operational life data in Table 2 were collected from reports by IRENA [7,8,9], which provide generic estimates for these parameters by technology. These reports also provide projections of future costs for renewable energy technologies. These data are presented in Table 3 and Figure 1, where it was assumed that costs fall linearly between the data points provided by IRENA and that costs remain constant beyond 2040 when the IRENA forecasts end (except for offshore wind, where the IRENA forecast continues to 2050).".format(country))

    document.add_paragraph('Country-specific capacity factors for solar PV, wind and hydropower were sourced from Renewables Ninja and the PLEXOS-World 2015 Model Dataset [3,10,11]. These sources provide hourly capacity factors for 2015 for solar PV and wind, and 15-year averages monthly capacity factors for hydropower, the average values of which are presented in Table 2. These data were also used to estimate capacity factors for 8 time slices used in the OSeMOSYS model (see detail in Annex 1). Capacity factors for other technologies were sourced from reports by IRENA [8,12], which provide generic estimates for each technology. The costs and efficiencies of power transmission and distribution were sourced from TEMBA reference case [23], which provides generic cost estimates and country-specific efficiencies which consider expected efficiency improvements in the future. Techno-economic data for refineries were sourced from the IEA Energy Technology Systems Analysis Programme (ETSAP) [15], which provides generic estimates of costs and performance parameters, while the refinery options modelled are based on the methods used in TEMBA [13].')

    document.add_heading('2.2 Fuel Data', 2)
    document.add_paragraph('The crude oil price is based on an international price forecast produced by the US Energy Information Administration (EIA), which runs to 2050 [16]. The price was increased by 10% for imported oil to reflect the cost of importation. The price of imported HFO and LFO were calculated by multiplying the oil price by 0.8 and 1.33 respectively, based on the methods used in TEMBA [13]. The prices of coal, natural gas and biomass were sourced from an IRENA report [8], which provides generic estimates for costs to 2030. Again, a linear rate of change was assumed between data points from IRENA, and the forecast was extended to 2040 using the rate of change between 2020 and 2030. Prices were then assumed constant after 2040. The cost of domestically-produced biomass was increased by 10% to estimate a cost of imported biomass.')

    document.add_heading('2.3 Emissions Factors and Domestic Reserves', 2)
    document.add_paragraph('Emissions factors were collected from the IPCC Emission Factor Database [17], which provides carbon emissions factors by fuel. Domestic renewable energy potentials for solar PV, CSP and wind were collected from an IRENA-KTH working paper [18], which provides estimates of potential yearly generation by country in Africa. Other renewable energy potentials were sourced from a regional report by IRENA [19] and the World Small Hydropower Development Report [20], which provide estimated potentials in MW by country. Estimated domestic fossil fuel reserves are from the websites of The World Bank and US EIA [21,22], which provide estimates of reserves by country. ')

    document.add_heading('2.4 Electricity Demand Data', 2)
    document.add_paragraph('The total electricity demand projection is based on data from the TEMBA Reference Scenario dataset [22], which provides yearly total demand estimates from 2015-2070 under a reference case scenario.')

    document.add_heading('3 Ethics Statement', 2)
    document.add_paragraph('Not applicable.')

    document.add_heading('4 CRediT Author Statement', 2)
    credit = document.add_paragraph('Lucy Allington: Data curation; Investigation; Methodology; Writing – original draft; Visualisation. Carla Cannone: Data curation; Investigation; Software; Formal analysis; Visualisation. Ioannis Pappis: Data curation; Investigation; Validation; Writing - Review & Editing. ')
    if not authors.empty:
        asplit = authors.split(',')
        for a in asplit:
            st = '{}: Investigation.'.format(a)
            credit.add_run(st)
    credit.add_run('Karla Cervantes Barron: Data Curation; Software; Visualisation. William Usher: Software; Supervision. Steve Pye: Supervision; Project Administration.  Mark Howells: Conceptualisation; Methodology; Writing – Review & Editing; Supervision. Constantinos Taliotis: Conceptualisation. Caroline Sundin: Conceptualisation. Vignesh Sridharan: Conceptualisation. Eunice Ramos: Conceptualisation. Maarten Brinkerink: Data curation. Paul Deane: Data Curation. Gustavo Moura: Data Curation. Arnaud Rouget: Conceptualisation. Andrii Gritsevskyi: Conceptualisation. David Wogan: Conceptualisation. Edito Barcelona: Conceptualisation. Holger Rogner: Conceptualisation.')

    document.add_heading('Acknowledgements', 2)
    document.add_paragraph('We would like to acknowledge data providers who helped make this, and future iterations possible, they include IEA, UNSTATS, APEC, IRENA, UCC, KTH, UFOP and others.')

    document.add_heading('Funding', 2)
    document.add_paragraph("As well as support in kind provided by the employers of the authors of this note, we also acknowledge core funding from the Climate Compatible Growth Program (#CCG) of the UK's Foreign Development and Commonwealth Office (FCDO). The views expressed in this paper do not necessarily reflect the UK government’s official policies. ")

    document.add_heading('Declaration of Competing Interests ', 2)
    document.add_paragraph('The authors declare that they have no known competing financial interests or personal relationships which have or could be perceived to have influenced the work reported in this article.')

    document.add_heading('References', 2)
    # Add list of references by type of country
    for para in input_doc.paragraphs:
        get_para_data(document, para)

    document.add_page_break()
    document.add_heading('Appendix A – Zero-Order Tier 2 OSeMOSYS Model ', 2)
    document.add_paragraph('The data described above were used to create a simple zero-order Tier 2 energy systems model. As it is open source and free an OSeMOSYS model is calibrated and run with three example scenarios. Note that these scenarios in no way represent development trajectories of the country. This model and its results are intended to act as an example of what can be produced using the data in this article and a starting point for further model development. ')

    document.add_paragraph('U4RIA are goals to improve energy modelling [2]. They are short for Ubuntu (meaning community focused), retrievability, reusability, repeatability, interoperability and auditability. The model moves to partially meet U4RIA goals in that: ')
    # style = 'List Bullet',
    document.add_paragraph('We develop examples of results that can be used by other research communities, including energy and transport, and to aid mitigation strategies.', style = 'List Bullet')
    document.add_paragraph('The illustrative analyses are retrievable, reusable, repeatable. ', style = 'List Bullet')
    document.add_paragraph('As data are defined, elements of interoperability are feasible.', style = 'List Bullet')
    document.add_paragraph('And by virtue of the above the analysis could be audited or verified (that is not to say that it is ‘accurate’ but simply reproducible).', style = 'List Bullet')

    document.add_paragraph('In the OSeMOSYS model, the electricity supply system is represented by importing and extraction technologies, conversion technologies, power plants, transmission and distribution network systems and final energy demands for the different available fuels considered. The Reference Energy System is shown below. The main modelling assumptions consist of power generation capacity per type of technology (centralized, decentralized), fuel prices, emissions, transmission and distribution network capacity and losses, and refineries, which are exogenous parameters into the model. Furthermore, the final energy demands which are exogenously entered into the model are disaggregated by fuel and sector. The data described in this article were used as input data to define these assumptions in the model.')
    document.add_paragraph('Figure 1: Reference Energy System ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/RefEnSyst.jpeg', width=Inches(6))

    document.add_heading('A1 Model Assumptions', 2)
    document.add_paragraph('Key assumptions used in model development are outlined below. ')

    document.add_heading('Supply-Side Assumptions ', 3)
    document.add_paragraph('The share of total demand that can be met by off-grid solar PV is constrained based on the optimal balance of on- and off-grid provision in a least cost scenario for 100% electricity access by 2030 from the Global Electrification Platform [25]. Additional technologies were modelled to represent utility-scale solar PV and onshore wind with storage capacity. Utility-scale PV with two-hour storage and onshore wind with half-hour storage were modelled, with the additional costs of storage estimated based on data from the NREL ATB 2020 Database, which provides cost projections for different durations of storage up to 2050 [26]. The maximum share of total demand that can be met by variable renewables is constrained as follows: utility-scale PV, onshore wind and utility-scale PV with storage are each permitted to meet up to 15% of demand; offshore wind can meet up to 10% of demand and onshore wind with storage can meet up to 25% of demand. This analysis is not intended to offer a detailed study of system flexibility; however these constraints are included to ensure the system is operational under high renewable shares. Biomass is permitted to meet up to 30% of electricity demand. Electricity imports and exports were modelled in a simplified manner whereby single import and exports technologies are constrained to import and export electricity in line with energy balance data [27]. ')

    document.add_heading('Demand-Side Assumptions ', 3)
    document.add_paragraph('Generic techno-economic data for demand-side technologies (cooking, heating and transport) were used [28,29]. The total electricity demand shown in Figure 3 was split by sector based on the proportions of demand in historical energy balance data [27]. In each sector, moderate and high energy efficiency technologies were modelled, with input activity ratios of 1 and output activity ratios of 1.15 and 1.3 respectively. This is a simplified way of allowing the model to invest in energy efficiency in each sector, with costs estimated based on the costs of electricity generation by a coal power plant in the model. In the Least Cost and Net Zero scenario (detailed in Section A2), there is a constraint on the speed at which fuel switching and energy efficiency investments can occur to better align results to reality. This is done by limiting the annual investment in electric vehicles, stoves, heating technologies and energy efficiency to 5% of the 2050 capacity.')

    document.add_heading('Time Representation and Discount Rate ', 3)
    document.add_paragraph('Within each model year, four seasons, each with two dayparts, are defined. Daypart 1 starts at 06:00 and finishes at 18:00, while daypart 2 starts at 18:00 and finishes at 06:00. The seasons are defined so that season 1 runs from December to February, season 2 runs from March to May, season 3 from June to August and season 4 from September to November. A discount rate of 10% is used. ')

    document.add_heading('A2 Scenario Definitions ', 2)
    document.add_paragraph('Three stylized scenarios are modelled: Fossil Future, Least Cost and Net Zero by 2050. These scenarios are defined in the table below. Nuclear power is not considered in any of these scenarios; however it can be added using the techno-economic data provided in the main article.')

    records3 = (
        ('Fossil Future ', 'No new investments in renewable or nuclear power generation, electric stoves and heating, electric transport or energy efficiency are permitted. '),
        ('Least Cost ', 'No new investment in nuclear power is permitted. Gradual investment constraints are applied to demand-side fuel-switching and energy efficiency, whereby only up to 5% of each technology’s capacity in 2050 can be invested in annually. No additional constraints are applied to find the cost-optimal solution. '),
        ('Net Zero by 2050 ', 'Domestic production and imports of fossil fuels and biomass gradually decline to 0 in 2050, beginning in 2021. No new investment in nuclear power is permitted. Gradual investment constraints are applied to demand-side fuel-switching and energy efficiency, whereby only up to 5% of each technology’s capacity in 2050 can be invested in annually from 2021-2039, rising to 10% from 2040-2050 to reflect greater ambition. '),
    )
    document.add_paragraph('Table A1: Definitions of the three model scenarios. ')
    table = document.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Scenario'
    hdr_cells[1].text = 'Definition'
    for s1, s2 in records3:
        row_cells = table.add_row().cells
        row_cells[0].text = s1
        row_cells[1].text = s2

    document.add_heading('A3 Scenario Results for {}'.format(country), 2)
    document.add_paragraph('The graphs below show selected results for the three modelled scenarios, including yearly electricity generation and supply capacity, fuel use in the transport sector and total annual carbon dioxide emissions for 2020-2050.')

    document.add_heading('A3.1 Electricity Generation Results', 3)
    # Add figure
    document.add_paragraph('Figure A1: Electricity Generation in the Fossil Future scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_FFAnnualElecProduction_{1}.jpg'.format(country,ic),
                         width=Inches(6))
    # Add figure
    document.add_paragraph('Figure A2: Electricity Generation in the Least Cost scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_LCv2AnnualElecProduction_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A3: Electricity Generation in the Net Zero by 2050  scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_NZv2AnnualElecProduction_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_heading('A3.2 Capacity Expansion Results', 3)

    document.add_paragraph('Figure A4: Installed capacity in the Fossil Future scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_FFTotalCapacityAnnual_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A5: Installed capacity in the Least Cost scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_LCv2TotalCapacityAnnual_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A6: Installed capacity in the Net Zero scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_NZv2TotalCapacityAnnual_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_heading('A3.3 Transport Results', 3)

    document.add_paragraph('Figure A7: Transport results for the Fossil Future scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_FFTransport_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A8: Transport results for the Least Cost scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_LCv2Transport_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A9: Transport results for the Net Zero scenario ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_NZv2Transport_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_heading('A3.4 Annual Carbon Dioxide Emissions Results', 3)

    document.add_paragraph('Figure A10: Annual Carbon Dioxide emissions in the Fossil Future scenario')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_FFAnnual_CO2_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A11: Annual Carbon Dioxide emissions in the Least Cost scenario  ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_LCv2Annual_CO2_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_paragraph('Figure A12: Annual Carbon Dioxide emissions in the Net Zero scenario  ')
    document.add_picture(r'C:/Users/KarlaC/CCG_Starter_Kits/figures/{0}/Fig_NZv2Annual_CO2_{1}.jpg'.format(country,ic),
                         width=Inches(6))

    document.add_heading('A4 Further Work ', 2)
    document.add_paragraph('These example results represent zero-order model and were generated using the clicSAND Interface [30] and OSeMOSYS code [31]. Those interested in further developing this work are directed to the dataset available on Zenodo [32] and guidance on model development using clicSAND and OSeMOSYS [33].  ')


    document.add_page_break()

    document.save(r'C:/Users/KarlaC/CCG_Starter_Kits/outputs/{0}/DiB_{1}.docx'.format(country,ic))

